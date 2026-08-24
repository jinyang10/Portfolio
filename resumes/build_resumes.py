#!/usr/bin/env python3
"""Generate targeted one-page resume variants from shared layout helpers."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.text.paragraph import Paragraph
from docx.text.run import Run

OUT_DIR = Path(__file__).resolve().parent
RIGHT_TAB = Twips(10773)
BULLET = "•"


def set_run_font(run: Run, *, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def set_paragraph_format(
    p: Paragraph,
    *,
    size: float = 10,
    after: int = 0,
    before: int = 0,
    align: WD_ALIGN_PARAGRAPH | None = None,
    left: int | None = None,
    hanging: int | None = None,
    right_tab: bool = True,
) -> None:
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(before)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = Twips(left)
    if hanging is not None:
        pf.first_line_indent = Twips(-hanging)
    if right_tab:
        pf.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    pPr = p._p.get_or_add_pPr()
    rpr = pPr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        pPr.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), str(int(size * 2)))


def add_hyperlink(paragraph: Paragraph, text: str, url: str, size: float) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_text(p: Paragraph, text: str, *, size: float, bold: bool = False, italic: bool = False) -> Run:
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return run


def new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Original Word margins: 510/567/510/624 twips
    section.top_margin = Twips(510)
    section.bottom_margin = Twips(510)
    section.left_margin = Twips(624)
    section.right_margin = Twips(567)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    return doc


def add_name(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, right_tab=False)
    add_text(p, name, size=14, bold=True)


def add_contact_line(doc: Document, items: list[tuple[str, str | None]]) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=9.5, after=4, align=WD_ALIGN_PARAGRAPH.CENTER, right_tab=False)
    for i, (label, url) in enumerate(items):
        if i:
            add_text(p, "  |  ", size=9.5)
        if url:
            add_hyperlink(p, label, url, 9.5)
        else:
            add_text(p, label, size=9.5)


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=11.5, before=8, after=1)
    run = add_text(p, title.upper(), size=11.5, bold=True)
    run.underline = True


def add_job_header(doc: Document, left: str, right: str = "") -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=10, after=0)
    add_text(p, left, size=10, bold=True)
    if right:
        add_text(p, "\t", size=10)
        add_text(p, right, size=10)


def add_italic_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=10, after=1)
    add_text(p, text, size=10, italic=True)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=10, after=0)
    if bold_prefix:
        add_text(p, bold_prefix, size=10, bold=True)
        add_text(p, text, size=10)
    else:
        add_text(p, text, size=10)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, size=10, after=0, left=220, hanging=180)
    add_text(p, f"{BULLET}  {text}", size=10)


def save(doc: Document, filename: str) -> Path:
    path = OUT_DIR / filename
    doc.save(path)
    return path


SHARED_CONTACT = [
    ("Toronto, ON", None),
    ("647-980-7007", "tel:6479807007"),
    ("sjy6@my.yorku.ca", "mailto:sjy6@my.yorku.ca"),
    ("GitHub", "https://github.com/jinyang10"),
    ("LinkedIn", "https://www.linkedin.com/in/jin-yang-aa7352301/"),
    ("Portfolio", "https://jinyang10.github.io/Portfolio/"),
]


def add_education(
    doc: Document,
    coursework: str,
    *,
    extra_line: str | None = None,
) -> None:
    add_section(doc, "Education")
    add_job_header(doc, "York University, Lassonde School of Engineering", "Toronto, ON")
    add_job_header(doc, "Bachelor of Engineering — Computer Engineering", "Expected June 2028")
    add_body(doc, "3.5/4.0", bold_prefix="GPA: ")
    add_body(doc, coursework, bold_prefix="Relevant Coursework: ")
    if extra_line:
        add_body(doc, extra_line)


def add_skills(doc: Document, rows: list[tuple[str, str]]) -> None:
    add_section(doc, "Technical Skills")
    for label, value in rows:
        add_body(doc, value, bold_prefix=f"{label}: ")


def build_big4() -> Path:
    doc = new_document()
    add_name(doc, "Jin Yang")
    add_contact_line(doc, SHARED_CONTACT)
    add_education(
        doc,
        "Software Design, Database Systems, Algorithms & Data Structures, Operating Systems",
    )
    add_skills(
        doc,
        [
            ("Languages", "Java, Python, C, SQL, TypeScript, Bash"),
            ("Frameworks & Libraries", "React, Next.js, Tailwind CSS, JDBC, JUnit"),
            ("Tools", "Git, Linux, PostgreSQL, SQLite, Node.js"),
        ],
    )

    add_section(doc, "Projects")

    add_job_header(doc, "MaternaDB — Midwifery Clinic Database", "Oct 2024 – Jan 2025")
    add_italic_line(doc, "Java, Swing, JDBC, SQL, SQLite, PostgreSQL")
    add_bullet(
        doc,
        "Designed a normalized relational schema (15+ entities in the original model; 11 core tables in the running app) covering midwives, pregnancies, appointments, notes, and tests.",
    )
    add_bullet(
        doc,
        "Built a Swing + JDBC client so a midwife can sign in, load a day’s appointments, review notes and lab results, add an observation, and prescribe a test.",
    )
    add_bullet(
        doc,
        "Ran the same queries against local SQLite or hosted PostgreSQL using environment-based credentials instead of hardcoded passwords.",
    )

    add_job_header(doc, "YU Lab Reservation System", "Jan 2026 – Apr 2026")
    add_italic_line(doc, "Java, Swing, JUnit, Randoop")
    add_bullet(
        doc,
        "Shipped a lab-equipment booking app with role-based users (student, faculty, researcher, guest, staff), reservations, payments, and approvals.",
    )
    add_bullet(
        doc,
        "Used Strategy, State, Observer, and Factory patterns for role-based pricing, reservation lifecycle (confirm, arrive, extend, cancel, no-show), and equipment monitoring.",
    )
    add_bullet(
        doc,
        "Added JUnit tests plus Randoop-generated regression suites around booking, payment, and registration services.",
    )

    add_job_header(doc, "Luma Health (in progress)", "May 2026 – Present")
    add_italic_line(doc, "Next.js, React, TypeScript, Tailwind CSS, shadcn/ui")
    add_bullet(
        doc,
        "Building a healthcare web app shell: responsive navigation, mobile drawer, theme toggle, and reusable department/doctor cards.",
    )
    add_bullet(
        doc,
        "Laid out patient/provider booking entry points (Home, Book Appointment, Sign In) as the base for role-based workflows.",
    )

    add_job_header(doc, "CoreShell", "Oct 2024 – Dec 2024")
    add_italic_line(doc, "C")
    add_bullet(
        doc,
        "Implemented a Unix-like shell with 8 commands, a PCB-based ready queue, and concurrent execution of up to 3 programs.",
    )
    add_bullet(
        doc,
        "Coded FCFS, Round Robin (quantum 2), SJF, and AGING schedulers, plus paging with page tables, faults, and victim replacement.",
    )

    add_section(doc, "Extracurricular Experience")
    add_job_header(doc, "York University Robotics Society", "Sep 2024 – Apr 2025")
    add_bullet(
        doc,
        "Prototyped mechanical parts in TinkerCAD for 3D printing and wrote sensor-to-motor control for autonomous movement and obstacle stop.",
    )
    add_job_header(doc, "Automated Plant Watering", "2025")
    add_italic_line(doc, "Java, Firmata4j, Arduino/Grove, JFreeChart")
    add_bullet(
        doc,
        "Built a working prototype that reads Grove soil-moisture data, drives a pump to a wetness threshold, logs multi-day samples, and plots moisture vs. time.",
    )
    return save(doc, "JinYang_Resume_Big4.docx")


def build_faang_swe() -> Path:
    doc = new_document()
    add_name(doc, "Jin Yang")
    add_contact_line(doc, SHARED_CONTACT)
    add_education(
        doc,
        "Algorithms & Data Structures, Operating Systems, Computer Systems, Database Systems, Software Design",
    )
    add_skills(
        doc,
        [
            ("Languages", "C, Java, Python, SQL, TypeScript"),
            ("Systems & Tools", "Linux, Git, Make, JDBC, PostgreSQL, SQLite, JUnit"),
        ],
    )

    add_section(doc, "Projects")

    add_job_header(doc, "CoreShell — OS Shell, Paging & Scheduling", "Oct 2024 – Dec 2024")
    add_italic_line(doc, "C")
    add_bullet(
        doc,
        "Wrote a modular shell in C (interpreter, kernel, PCB, CPU, shell memory) that runs scripts and up to 3 programs via exec.",
    )
    add_bullet(
        doc,
        "Implemented FCFS, Round Robin (quantum 2), SJF, and AGING on a ready queue, with PCBs tracking PID, PC, and page-table state.",
    )
    add_bullet(
        doc,
        "Added demand paging: 3-line pages, per-process page tables, lazy loads into a frame store, page-fault handling, and victim replacement.",
    )

    add_job_header(doc, "YU Lab Reservation System", "Jan 2026 – Apr 2026")
    add_italic_line(doc, "Java, JUnit, Randoop")
    add_bullet(
        doc,
        "Designed an object-oriented booking system with role-specific users, equipment inventory, payments, and reservation state transitions.",
    )
    add_bullet(
        doc,
        "Separated pricing, payment, and lifecycle behind Strategy/State/Factory so new user types and payment methods extend without rewriting booking logic.",
    )
    add_bullet(
        doc,
        "Covered booking, payment, and registration paths with hand-written JUnit tests and Randoop regression suites.",
    )

    add_job_header(doc, "MaternaDB — Relational Clinic System", "Oct 2024 – Jan 2025")
    add_italic_line(doc, "Java, JDBC, SQL, SQLite, PostgreSQL")
    add_bullet(
        doc,
        "Modeled midwives, couples, pregnancies, appointments, notes, and tests with primary/foreign keys and a 3NF-style schema.",
    )
    add_bullet(
        doc,
        "Implemented parameterized JDBC queries (login, date lookup, multi-table appointment views, inserts for notes and prescribed tests).",
    )
    add_bullet(
        doc,
        "Kept the same SQL working on SQLite and PostgreSQL so the app is runnable without a licensed DB2 instance.",
    )

    add_job_header(doc, "Luma Health (in progress)", "May 2026 – Present")
    add_italic_line(doc, "TypeScript, React, Next.js")
    add_bullet(
        doc,
        "Componentized a Next.js/TypeScript UI (header, mobile sheet, theming, department/doctor cards) as the start of a role-based healthcare app.",
    )

    add_section(doc, "Additional")
    add_job_header(doc, "York University Robotics Society", "Sep 2024 – Apr 2025")
    add_bullet(
        doc,
        "Integrated sensor input and motor control for autonomous movement and obstacle detection on a student robot.",
    )
    return save(doc, "JinYang_Resume_FAANG_SWE.docx")


def build_hardware() -> Path:
    doc = new_document()
    add_name(doc, "Jin Yang")
    add_contact_line(doc, SHARED_CONTACT)
    add_education(
        doc,
        "Digital Logic Design, Electronic Circuits & Devices, Computer Systems, Operating Systems, Algorithms & Data Structures",
    )
    add_skills(
        doc,
        [
            ("Languages", "C, Verilog, Java, Python, MATLAB, Bash"),
            ("Hardware & Tools", "FPGA (DE10-Lite), VGA, Arduino/Grove, LTspice, PSpice, Linux, Git"),
        ],
    )

    add_section(doc, "Projects")

    add_job_header(doc, "Snake Game on FPGA", "Nov 2025 – Dec 2025")
    add_italic_line(doc, "Verilog, FPGA (DE10-Lite), VGA")
    add_bullet(
        doc,
        "Implemented real-time Snake on a DE10-Lite with 640×480 VGA output, including sync/timing logic for stable frames.",
    )
    add_bullet(
        doc,
        "Used finite-state machines for game state, movement, and rendering updates in hardware.",
    )
    add_bullet(
        doc,
        "Integrated accelerometer input for low-latency directional control instead of buttons-only play.",
    )

    add_job_header(doc, "Automated Plant Watering System", "2025")
    add_italic_line(doc, "Java, Firmata4j, Arduino/Grove, JFreeChart")
    add_bullet(
        doc,
        "Wrote a Java controller that reads Grove soil-moisture samples over Firmata and drives a pump until the soil hits a wetness threshold.",
    )
    add_bullet(
        doc,
        "Ran the loop across multiple days, logged moisture vs. time, and plotted the series with JFreeChart on shutdown.",
    )

    add_job_header(doc, "CoreShell — Embedded-style OS Runtime in C", "Oct 2024 – Dec 2024")
    add_italic_line(doc, "C, Linux")
    add_bullet(
        doc,
        "Built a C shell with process control blocks, a ready queue, and concurrent execution of up to 3 programs.",
    )
    add_bullet(
        doc,
        "Implemented four CPU schedulers and a paging layer (page tables, faults, frame-store victim replacement) on Linux.",
    )

    add_section(doc, "Experience")
    add_job_header(doc, "York University Robotics Society", "Sep 2024 – Apr 2025")
    add_italic_line(doc, "TinkerCAD, sensors, motor control")
    add_bullet(
        doc,
        "Designed and 3D-printed mechanical parts in TinkerCAD with an eye toward weight distribution and structural integrity.",
    )
    add_bullet(
        doc,
        "Integrated sensor input and motor-control logic so the robot could move autonomously and stop on obstacle detect.",
    )
    return save(doc, "JinYang_Resume_Hardware_Embedded.docx")


def main() -> None:
    paths = [build_big4(), build_faang_swe(), build_hardware()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
